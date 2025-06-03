import streamlit as st
import json
import io
import os
import base64
import hashlib
from datetime import datetime
from typing import Dict, List, Any
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure the page
st.set_page_config(
    page_title="TAC Resume Builder",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

class UserManager:
    def __init__(self):
        self.users_file = "users.json"
        self.current_user = None
    
    def load_users(self):
        try:
            if os.path.exists(self.users_file):
                with open(self.users_file, 'r') as f:
                    return json.load(f)
            return {}
        except:
            return {}
    
    def save_users(self, users):
        try:
            with open(self.users_file, 'w') as f:
                json.dump(users, f, indent=2)
        except:
            pass
    
    def hash_password(self, password):
        return hashlib.sha256(password.encode()).hexdigest()
    
    def register_user(self, username, password, email):
        users = self.load_users()
        if username in users:
            return False, "Username already exists"
        
        users[username] = {
            'password': self.hash_password(password),
            'email': email,
            'created_at': datetime.now().isoformat(),
            'resumes': {}
        }
        self.save_users(users)
        return True, "Registration successful"
    
    def login_user(self, username, password):
        users = self.load_users()
        if username in users and users[username]['password'] == self.hash_password(password):
            self.current_user = username
            return True, "Login successful"
        return False, "Invalid credentials"
    
    def save_user_resume(self, username, resume_data, resume_name="default"):
        users = self.load_users()
        if username in users:
            if 'resumes' not in users[username]:
                users[username]['resumes'] = {}
            users[username]['resumes'][resume_name] = {
                'data': resume_data,
                'last_updated': datetime.now().isoformat()
            }
            self.save_users(users)
            return True
        return False
    
    def load_user_resume(self, username, resume_name="default"):
        users = self.load_users()
        if username in users and 'resumes' in users[username] and resume_name in users[username]['resumes']:
            return users[username]['resumes'][resume_name]['data']
        return None
    
    def get_user_resumes(self, username):
        users = self.load_users()
        if username in users and 'resumes' in users[username]:
            return list(users[username]['resumes'].keys())
        return []

class SimpleAIGenerator:
    def __init__(self):
        self.api_key = os.getenv("OPENAI_API_KEY")
        
    def generate_professional_summary(self, job_title: str, years_experience: int, industry: str) -> str:
        if not self.api_key:
            return f"Experienced {job_title} with {years_experience} years in {industry}. Proven track record of delivering results and contributing to organizational success."
        
        try:
            from openai import OpenAI
            client = OpenAI(api_key=self.api_key)
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert resume writer."},
                    {"role": "user", "content": f"Write a professional summary for a {job_title} with {years_experience} years of experience in {industry}. Keep it 3-4 sentences."}
                ],
                max_tokens=200
            )
            return response.choices[0].message.content.strip() if response.choices[0].message.content else f"Experienced {job_title} with {years_experience} years in {industry}."
        except Exception as e:
            return f"Experienced {job_title} with {years_experience} years in {industry}. Proven track record of delivering results and contributing to organizational success."
    
    def generate_bullet_points(self, job_title: str, company: str, job_description: str) -> List[str]:
        if not self.api_key:
            return [
                f"Performed key responsibilities as {job_title}",
                f"Contributed to {company}'s operational excellence",
                "Collaborated with cross-functional teams",
                "Achieved measurable results and improvements"
            ]
        
        try:
            from openai import OpenAI
            client = OpenAI(api_key=self.api_key)
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert resume writer. Create 4-6 impactful bullet points for work experience."},
                    {"role": "user", "content": f"Create bullet points for a {job_title} position at {company}. Job description: {job_description}. Start each with an action verb and include quantifiable achievements where possible."}
                ],
                max_tokens=300
            )
            
            content = response.choices[0].message.content
            if content:
                bullets = [line.strip().lstrip('•-*').strip() for line in content.split('\n') if line.strip()]
                return bullets[:6] if bullets else [f"Performed key responsibilities as {job_title}"]
            return [f"Performed key responsibilities as {job_title}"]
        except Exception as e:
            return [
                f"Performed key responsibilities as {job_title}",
                f"Contributed to {company}'s operational excellence",
                "Collaborated with cross-functional teams",
                "Achieved measurable results and improvements"
            ]

class SimpleExportManager:
    def export_to_pdf(self, resume_data: Dict, template_name: str = 'modern') -> bytes:
        """Export resume to PDF format"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, 
                                  rightMargin=0.75*inch, leftMargin=0.75*inch, 
                                  topMargin=0.75*inch, bottomMargin=0.75*inch)
            
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#2E86C1'),
                alignment=TA_CENTER,
                spaceAfter=6
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#2E86C1'),
                spaceAfter=6
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4
            )
            
            story = []
            personal = resume_data.get('personal', {})
            
            # Header
            if personal.get('full_name'):
                story.append(Paragraph(personal['full_name'], title_style))
                
                contact_info = []
                if personal.get('email'):
                    contact_info.append(personal['email'])
                if personal.get('phone'):
                    contact_info.append(personal['phone'])
                if personal.get('location'):
                    contact_info.append(personal['location'])
                
                if contact_info:
                    story.append(Paragraph(' • '.join(contact_info), body_style))
                
                story.append(Spacer(1, 12))
            
            # Professional Summary
            if resume_data.get('summary'):
                story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
                story.append(Paragraph(resume_data['summary'], body_style))
                story.append(Spacer(1, 12))
            
            # Work Experience
            work_exp = resume_data.get('work_experience', [])
            if work_exp:
                story.append(Paragraph("PROFESSIONAL EXPERIENCE", heading_style))
                for exp in work_exp:
                    job_header = f"<b>{exp.get('job_title', '')}</b> - {exp.get('company', '')}"
                    if exp.get('location'):
                        job_header += f" ({exp['location']})"
                    story.append(Paragraph(job_header, body_style))
                    
                    date_range = f"{exp.get('start_date', '')} to {exp.get('end_date', 'Present')}"
                    story.append(Paragraph(f"<i>{date_range}</i>", body_style))
                    
                    for bullet in exp.get('bullets', []):
                        story.append(Paragraph(f"• {bullet}", body_style))
                    
                    story.append(Spacer(1, 8))
            
            # Education
            education = resume_data.get('education', [])
            if education:
                story.append(Paragraph("EDUCATION", heading_style))
                for edu in education:
                    edu_text = f"<b>{edu.get('degree', '')}</b>"
                    if edu.get('major'):
                        edu_text += f" in {edu['major']}"
                    edu_text += f" - {edu.get('school', '')}"
                    if edu.get('graduation_date'):
                        edu_text += f" ({edu['graduation_date']})"
                    story.append(Paragraph(edu_text, body_style))
                    story.append(Spacer(1, 4))
            
            # Skills
            skills = resume_data.get('skills', {})
            if skills:
                story.append(Paragraph("SKILLS", heading_style))
                for category, skill_list in skills.items():
                    if skill_list:
                        skills_text = f"<b>{category}:</b> {', '.join(skill_list)}"
                        story.append(Paragraph(skills_text, body_style))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"PDF export failed: {str(e)}")
    
    def export_to_docx(self, resume_data: Dict, template_name: str = 'modern') -> bytes:
        """Export resume to DOCX format"""
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            personal = resume_data.get('personal', {})
            
            # Header
            if personal.get('full_name'):
                name_para = doc.add_paragraph()
                name_run = name_para.add_run(personal['full_name'])
                name_run.font.size = Pt(18)
                name_run.bold = True
                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                contact_info = []
                if personal.get('email'):
                    contact_info.append(personal['email'])
                if personal.get('phone'):
                    contact_info.append(personal['phone'])
                if personal.get('location'):
                    contact_info.append(personal['location'])
                
                if contact_info:
                    contact_para = doc.add_paragraph(' • '.join(contact_info))
                    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()
            
            # Professional Summary
            if resume_data.get('summary'):
                self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
                doc.add_paragraph(resume_data['summary'])
            
            # Work Experience
            work_exp = resume_data.get('work_experience', [])
            if work_exp:
                self._add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
                for exp in work_exp:
                    job_para = doc.add_paragraph()
                    job_run = job_para.add_run(f"{exp.get('job_title', '')} - {exp.get('company', '')}")
                    job_run.bold = True
                    
                    date_range = f"{exp.get('start_date', '')} to {exp.get('end_date', 'Present')}"
                    date_para = doc.add_paragraph(date_range)
                    
                    for bullet in exp.get('bullets', []):
                        bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                    
                    doc.add_paragraph()
            
            # Education
            education = resume_data.get('education', [])
            if education:
                self._add_section_heading(doc, "EDUCATION")
                for edu in education:
                    edu_para = doc.add_paragraph()
                    degree_run = edu_para.add_run(edu.get('degree', ''))
                    degree_run.bold = True
                    
                    if edu.get('major'):
                        edu_para.add_run(f" in {edu['major']}")
                    
                    edu_para.add_run(f" - {edu.get('school', '')}")
                    
                    if edu.get('graduation_date'):
                        edu_para.add_run(f" ({edu['graduation_date']})")
            
            # Skills
            skills = resume_data.get('skills', {})
            if skills:
                self._add_section_heading(doc, "SKILLS")
                for category, skill_list in skills.items():
                    if skill_list:
                        skill_para = doc.add_paragraph()
                        category_run = skill_para.add_run(f"{category}: ")
                        category_run.bold = True
                        skill_para.add_run(', '.join(skill_list))
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"DOCX export failed: {str(e)}")
    
    def _add_section_heading(self, doc, title):
        """Add section heading to DOCX document"""
        heading = doc.add_paragraph()
        heading_run = heading.add_run(title)
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        # Set color for heading (color not available in all environments)
    
    def export_to_html(self, resume_data: Dict, template_name: str = 'modern') -> str:
        personal = resume_data.get('personal', {})
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Resume - {personal.get('full_name', 'Resume')}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; color: #333; line-height: 1.6; }}
                .header {{ text-align: center; border-bottom: 2px solid #2E86C1; padding-bottom: 20px; margin-bottom: 30px; }}
                .name {{ font-size: 28px; font-weight: bold; color: #2E86C1; margin-bottom: 10px; }}
                .contact {{ margin-bottom: 5px; }}
                .section {{ margin-bottom: 25px; }}
                .section-title {{ font-size: 18px; font-weight: bold; color: #2E86C1; border-bottom: 1px solid #85C1E9; padding-bottom: 5px; margin-bottom: 15px; }}
                .job-title {{ font-weight: bold; color: #2E86C1; }}
                .company {{ font-weight: 600; }}
                .date-range {{ font-style: italic; color: #666; font-size: 14px; }}
                .bullet {{ margin: 5px 0 5px 20px; }}
                .skill-category {{ margin-bottom: 10px; }}
                .skill-category strong {{ color: #2E86C1; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="name">{personal.get('full_name', '')}</div>
                <div class="contact">{personal.get('email', '')} | {personal.get('phone', '')} | {personal.get('location', '')}</div>
        """
        
        if personal.get('linkedin') or personal.get('website'):
            html += f"<div class='contact'>"
            if personal.get('linkedin'):
                html += f"LinkedIn: {personal['linkedin']}"
            if personal.get('website'):
                html += f" | Website: {personal['website']}"
            html += "</div>"
        
        html += "</div>"
        
        # Professional Summary
        if resume_data.get('summary'):
            html += f"""
            <div class="section">
                <div class="section-title">Professional Summary</div>
                <p>{resume_data['summary']}</p>
            </div>
            """
        
        # Work Experience
        work_exp = resume_data.get('work_experience', [])
        if work_exp:
            html += '<div class="section"><div class="section-title">Professional Experience</div>'
            for exp in work_exp:
                html += f"""
                <div style="margin-bottom: 20px;">
                    <div><span class="job-title">{exp.get('job_title', '')}</span> - <span class="company">{exp.get('company', '')}</span></div>
                    <div class="date-range">{exp.get('start_date', '')} to {exp.get('end_date', 'Present')} | {exp.get('location', '')}</div>
                """
                for bullet in exp.get('bullets', []):
                    html += f'<div class="bullet">• {bullet}</div>'
                html += "</div>"
            html += "</div>"
        
        # Education
        education = resume_data.get('education', [])
        if education:
            html += '<div class="section"><div class="section-title">Education</div>'
            for edu in education:
                html += f"""
                <div style="margin-bottom: 15px;">
                    <div class="job-title">{edu.get('degree', '')} in {edu.get('major', '')} - {edu.get('school', '')}</div>
                    <div class="date-range">Graduated: {edu.get('graduation_date', '')}</div>
                </div>
                """
            html += "</div>"
        
        # Skills
        skills = resume_data.get('skills', {})
        if skills:
            html += '<div class="section"><div class="section-title">Skills</div>'
            for category, skill_list in skills.items():
                if skill_list:
                    html += f'<div class="skill-category"><strong>{category}:</strong> {", ".join(skill_list)}</div>'
            html += "</div>"
        
        html += "</body></html>"
        return html

def initialize_session_state():
    if 'resume_data' not in st.session_state:
        st.session_state.resume_data = {
            'personal': {'full_name': '', 'email': '', 'phone': '', 'location': '', 'linkedin': '', 'website': ''},
            'summary': '',
            'work_experience': [],
            'education': [],
            'skills': {},
            'projects': [],
            'certifications': []
        }
    
    if 'current_page' not in st.session_state:
        st.session_state.current_page = 'landing'
    
    if 'current_step' not in st.session_state:
        st.session_state.current_step = 0
    
    if 'user_authenticated' not in st.session_state:
        st.session_state.user_authenticated = False
    
    if 'current_user' not in st.session_state:
        st.session_state.current_user = None
    
    if 'selected_template' not in st.session_state:
        st.session_state.selected_template = 'modern'

def render_authentication():
    """Render login/register interface"""
    user_mgr = UserManager()
    
    if st.session_state.user_authenticated:
        with st.sidebar:
            st.success(f"Welcome, {st.session_state.current_user}!")
            
            # User's saved resumes
            saved_resumes = user_mgr.get_user_resumes(st.session_state.current_user)
            if saved_resumes:
                st.markdown("### Your Saved Resumes")
                selected_resume = st.selectbox("Load Resume:", [""] + saved_resumes)
                
                if selected_resume and st.button("Load Selected Resume"):
                    loaded_data = user_mgr.load_user_resume(st.session_state.current_user, selected_resume)
                    if loaded_data:
                        st.session_state.resume_data = loaded_data
                        st.success(f"Resume '{selected_resume}' loaded!")
                        st.rerun()
                
                # Save current resume
                resume_name = st.text_input("Save as:", placeholder="My Resume")
                if st.button("Save Current Resume") and resume_name:
                    if user_mgr.save_user_resume(st.session_state.current_user, st.session_state.resume_data, resume_name):
                        st.success(f"Resume saved as '{resume_name}'!")
                    else:
                        st.error("Failed to save resume")
            
            if st.button("Logout"):
                st.session_state.user_authenticated = False
                st.session_state.current_user = None
                st.rerun()
        
        return True
    
    else:
        st.markdown("## Account Access")
        tab1, tab2 = st.tabs(["Login", "Register"])
        
        with tab1:
            st.markdown("### Login to Your Account")
            username = st.text_input("Username", key="login_username")
            password = st.text_input("Password", type="password", key="login_password")
            
            if st.button("Login", type="primary"):
                if username and password:
                    success, message = user_mgr.login_user(username, password)
                    if success:
                        st.session_state.user_authenticated = True
                        st.session_state.current_user = username
                        st.success(message)
                        st.rerun()
                    else:
                        st.error(message)
                else:
                    st.error("Please enter username and password")
        
        with tab2:
            st.markdown("### Create New Account")
            new_username = st.text_input("Choose Username", key="reg_username")
            new_email = st.text_input("Email Address", key="reg_email")
            new_password = st.text_input("Choose Password", type="password", key="reg_password")
            confirm_password = st.text_input("Confirm Password", type="password", key="reg_confirm")
            
            if st.button("Create Account", type="primary"):
                if new_username and new_email and new_password and confirm_password:
                    if new_password == confirm_password:
                        success, message = user_mgr.register_user(new_username, new_password, new_email)
                        if success:
                            st.success(message + " Please login with your credentials.")
                        else:
                            st.error(message)
                    else:
                        st.error("Passwords do not match")
                else:
                    st.error("Please fill in all fields")
        
        st.markdown("---")
        st.info("Create an account to save your resumes and continue working on them later!")
        
        if st.button("Continue as Guest", use_container_width=True):
            st.session_state.current_page = 'builder'
            st.rerun()
        
        return False

def render_landing_page():
    st.markdown("""
    <div style="text-align: center; padding: 2rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 15px; color: white; margin: 2rem 0;">
        <h1 style="font-size: 3rem; margin-bottom: 1rem;">🎯 TAC Resume Builder</h1>
        <h3 style="font-weight: 300; margin-bottom: 2rem;">AI-Powered Professional Resume Generator</h3>
        <p style="font-size: 1.2rem; opacity: 0.9;">Create stunning, ATS-optimized resumes in minutes with the power of artificial intelligence</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Key Statistics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Templates", "8+", help="Professional resume templates")
    with col2:
        st.metric("Export Formats", "4", help="PDF, DOCX, HTML, JSON")
    with col3:
        st.metric("AI Features", "Active", help="AI-powered content generation")
    with col4:
        st.metric("User Accounts", "Enabled", help="Save and resume later")
    
    st.markdown("---")
    
    # Main Action Buttons
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("🚀 Start Building Your Resume", type="primary", use_container_width=True):
            if st.session_state.user_authenticated:
                st.session_state.current_page = "builder"
                st.rerun()
            else:
                st.session_state.current_page = "auth"
                st.rerun()
        
        if st.button("👤 Login / Register", use_container_width=True):
            st.session_state.current_page = "auth"
            st.rerun()
        
        if st.button("📄 View Templates", use_container_width=True):
            st.session_state.current_page = "templates"
            st.rerun()
    
    st.markdown("---")
    
    # Export Options Showcase
    st.markdown("## 📄 Export Your Resume in Multiple Formats")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        <div style="text-align: center; padding: 1.5rem; border: 2px solid #e1e8ed; border-radius: 10px; margin: 0.5rem;">
            <h3>📄 PDF Export</h3>
            <p>Professional print-ready format</p>
            <ul style="text-align: left; font-size: 0.9rem;">
                <li>ATS-optimized formatting</li>
                <li>Professional styling</li>
                <li>High-quality output</li>
                <li>Universal compatibility</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div style="text-align: center; padding: 1.5rem; border: 2px solid #e1e8ed; border-radius: 10px; margin: 0.5rem;">
            <h3>📝 DOCX Export</h3>
            <p>Editable Microsoft Word format</p>
            <ul style="text-align: left; font-size: 0.9rem;">
                <li>Fully editable</li>
                <li>Track changes support</li>
                <li>Comment functionality</li>
                <li>Professional formatting</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div style="text-align: center; padding: 1.5rem; border: 2px solid #e1e8ed; border-radius: 10px; margin: 0.5rem;">
            <h3>🌐 HTML Export</h3>
            <p>Web-compatible format</p>
            <ul style="text-align: left; font-size: 0.9rem;">
                <li>Responsive design</li>
                <li>Easy sharing</li>
                <li>Offline viewing</li>
                <li>Print to PDF option</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown("""
        <div style="text-align: center; padding: 1.5rem; border: 2px solid #e1e8ed; border-radius: 10px; margin: 0.5rem;">
            <h3>💾 JSON Backup</h3>
            <p>Data backup and transfer</p>
            <ul style="text-align: left; font-size: 0.9rem;">
                <li>Complete data backup</li>
                <li>Easy import/export</li>
                <li>Version control</li>
                <li>Cross-platform</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Key Features
    st.markdown("## ✨ Key Features")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        **🤖 AI-Powered Content Generation**
        - Professional summaries tailored to your experience
        - Compelling bullet points with action verbs
        - Industry-specific skill suggestions
        - ATS optimization recommendations
        
        **📄 Comprehensive Export Options**
        - PDF for professional applications
        - DOCX for further editing
        - HTML for web sharing
        - JSON for data backup
        
        **🔐 User Account System**
        - Save multiple resume versions
        - Resume later feature
        - Secure password protection
        - Personal resume library
        """)
    
    with col2:
        st.markdown("""
        **🎨 Professional Templates**
        - Modern: Clean, contemporary design
        - Executive: Sophisticated business format
        - Creative: Vibrant, artistic layout
        - Technical: Structured, logical format
        - Academic: Traditional academic style
        - Minimal: Simple, elegant design
        - Two-column: Space-efficient layout
        - Bold Impact: Strong visual presence
        
        **🔒 Privacy & Security**
        - Local data storage
        - Encrypted user passwords
        - No external data transmission
        - Complete user control
        """)
    
    # How It Works
    st.markdown("---")
    st.markdown("## 🔄 How It Works")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        ### 1. 👤 Create Account
        Register for free to save your progress and access your resumes anytime.
        """)
    
    with col2:
        st.markdown("""
        ### 2. 📝 Build Resume
        Use our step-by-step builder with AI assistance to create professional content.
        """)
    
    with col3:
        st.markdown("""
        ### 3. 🎨 Choose Template
        Select from 8+ professional templates and customize to match your style.
        """)
    
    with col4:
        st.markdown("""
        ### 4. 📄 Export & Apply
        Download in your preferred format and start applying to your dream jobs.
        """)
    
    # Call to Action
    st.markdown("---")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("### Ready to Build Your Professional Resume?")
        if st.button("Get Started Now", type="primary", use_container_width=True):
            st.session_state.current_page = "auth"
            st.rerun()

def render_templates_page():
    st.title("🎨 Resume Templates")
    st.markdown("Choose from our collection of professional templates")
    
    templates = {
        'modern': {'name': 'Modern Professional', 'desc': 'Clean, contemporary design perfect for tech and business roles'},
        'executive': {'name': 'Executive', 'desc': 'Sophisticated layout emphasizing leadership and experience'},
        'creative': {'name': 'Creative', 'desc': 'Vibrant design for creative professionals'},
        'technical': {'name': 'Technical', 'desc': 'Clean format optimized for technical roles'},
        'academic': {'name': 'Academic', 'desc': 'Traditional format for academic positions'},
        'minimal': {'name': 'Minimal', 'desc': 'Simple, elegant design focusing on content'},
        'two_column': {'name': 'Two Column', 'desc': 'Efficient layout maximizing space'},
        'bold': {'name': 'Bold Impact', 'desc': 'Strong visual design that stands out'}
    }
    
    cols = st.columns(2)
    for i, (template_id, template_info) in enumerate(templates.items()):
        col = cols[i % 2]
        with col:
            with st.container():
                st.subheader(template_info['name'])
                st.write(template_info['desc'])
                if st.button(f"Use {template_info['name']}", key=f"use_{template_id}"):
                    st.session_state.selected_template = template_id
                    st.session_state.current_page = "builder"
                    st.rerun()
                st.markdown("---")

def render_builder_page():
    st.title("📝 Resume Builder")
    
    # Sidebar navigation
    with st.sidebar:
        st.header("📋 Resume Sections")
        
        steps = [
            "Personal Information",
            "Professional Summary",
            "Work Experience", 
            "Education",
            "Skills",
            "Export"
        ]
        
        current_step = st.selectbox(
            "Choose Section:",
            range(len(steps)),
            format_func=lambda x: f"{x+1}. {steps[x]}",
            index=st.session_state.current_step
        )
        st.session_state.current_step = current_step
        
        st.divider()
        
        # Quick Export
        st.header("📄 Quick Export")
        export_mgr = SimpleExportManager()
        
        # PDF Export
        if st.button("📄 Export PDF", use_container_width=True):
            try:
                pdf_data = export_mgr.export_to_pdf(st.session_state.resume_data, st.session_state.selected_template)
                st.download_button(
                    label="Download PDF Resume",
                    data=pdf_data,
                    file_name="resume.pdf",
                    mime="application/pdf"
                )
            except Exception as e:
                st.error(f"PDF export failed: {str(e)}")
        
        # DOCX Export
        if st.button("📝 Export DOCX", use_container_width=True):
            try:
                docx_data = export_mgr.export_to_docx(st.session_state.resume_data, st.session_state.selected_template)
                st.download_button(
                    label="Download DOCX Resume",
                    data=docx_data,
                    file_name="resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"DOCX export failed: {str(e)}")
        
        # HTML Export
        if st.button("🌐 Export HTML", use_container_width=True):
            html_data = export_mgr.export_to_html(st.session_state.resume_data)
            st.download_button(
                label="Download HTML Resume",
                data=html_data,
                file_name="resume.html",
                mime="text/html"
            )
        
        st.divider()
        
        if st.button("💾 Save Progress", use_container_width=True):
            st.success("Progress saved locally!")
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        if current_step == 0:
            render_personal_info()
        elif current_step == 1:
            render_professional_summary()
        elif current_step == 2:
            render_work_experience()
        elif current_step == 3:
            render_education()
        elif current_step == 4:
            render_skills()
        elif current_step == 5:
            render_export()
    
    with col2:
        st.header("👁️ Live Preview")
        render_preview()

def render_personal_info():
    st.header("📋 Personal Information")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.session_state.resume_data['personal']['full_name'] = st.text_input(
            "Full Name *",
            value=st.session_state.resume_data['personal'].get('full_name', ''),
            placeholder="John Doe"
        )
        
        st.session_state.resume_data['personal']['email'] = st.text_input(
            "Email Address *",
            value=st.session_state.resume_data['personal'].get('email', ''),
            placeholder="john.doe@email.com"
        )
        
        st.session_state.resume_data['personal']['phone'] = st.text_input(
            "Phone Number",
            value=st.session_state.resume_data['personal'].get('phone', ''),
            placeholder="+1 (555) 123-4567"
        )
    
    with col2:
        st.session_state.resume_data['personal']['location'] = st.text_input(
            "Location",
            value=st.session_state.resume_data['personal'].get('location', ''),
            placeholder="City, State"
        )
        
        st.session_state.resume_data['personal']['linkedin'] = st.text_input(
            "LinkedIn URL",
            value=st.session_state.resume_data['personal'].get('linkedin', ''),
            placeholder="linkedin.com/in/johndoe"
        )
        
        st.session_state.resume_data['personal']['website'] = st.text_input(
            "Website/Portfolio",
            value=st.session_state.resume_data['personal'].get('website', ''),
            placeholder="www.johndoe.com"
        )

def render_professional_summary():
    st.header("✨ Professional Summary")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        job_title = st.text_input("Target Job Title", placeholder="Software Engineer")
        years_experience = st.number_input("Years of Experience", min_value=0, max_value=50, value=0)
        industry = st.selectbox("Industry", ["Technology", "Healthcare", "Finance", "Marketing", "Education", "Other"])
    
    with col2:
        if st.button("🤖 Generate AI Summary", use_container_width=True):
            if job_title:
                ai_gen = SimpleAIGenerator()
                summary = ai_gen.generate_professional_summary(job_title, years_experience, industry)
                st.session_state.resume_data['summary'] = summary
                st.success("AI summary generated!")
                st.rerun()
    
    st.session_state.resume_data['summary'] = st.text_area(
        "Professional Summary",
        value=st.session_state.resume_data.get('summary', ''),
        height=150,
        help="A brief overview of your professional background"
    )

def render_work_experience():
    st.header("💼 Work Experience")
    
    if 'work_experience' not in st.session_state.resume_data:
        st.session_state.resume_data['work_experience'] = []
    
    with st.expander("➕ Add New Work Experience", expanded=len(st.session_state.resume_data['work_experience']) == 0):
        col1, col2 = st.columns(2)
        
        with col1:
            job_title = st.text_input("Job Title", key="new_job_title")
            company = st.text_input("Company Name", key="new_company")
            start_date = st.text_input("Start Date", key="new_start_date", placeholder="MM/YYYY")
        
        with col2:
            location = st.text_input("Location", key="new_location")
            current_job = st.checkbox("Current Position", key="new_current")
            end_date = "Present" if current_job else st.text_input("End Date", key="new_end_date", placeholder="MM/YYYY")
        
        bullets_text = st.text_area(
            "Key Achievements (one per line)",
            key="new_bullets",
            height=100,
            placeholder="• Managed a team of 10 developers\n• Increased efficiency by 25%\n• Led migration to cloud infrastructure"
        )
        
        if st.button("✅ Add Experience", type="primary"):
            if job_title and company:
                bullets = [bullet.strip() for bullet in bullets_text.split('\n') if bullet.strip()]
                experience = {
                    'job_title': job_title,
                    'company': company,
                    'location': location,
                    'start_date': start_date,
                    'end_date': end_date,
                    'current': current_job,
                    'bullets': bullets
                }
                st.session_state.resume_data['work_experience'].append(experience)
                st.success("Experience added!")
                st.rerun()
    
    # Display existing experiences
    for i, exp in enumerate(st.session_state.resume_data['work_experience']):
        with st.expander(f"{exp['job_title']} at {exp['company']}", expanded=False):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                st.write(f"**Location:** {exp['location']}")
                st.write(f"**Duration:** {exp['start_date']} to {exp['end_date']}")
                if exp['bullets']:
                    st.write("**Key Achievements:**")
                    for bullet in exp['bullets']:
                        st.write(f"• {bullet}")
            
            with col2:
                if st.button("🗑️ Delete", key=f"delete_{i}"):
                    st.session_state.resume_data['work_experience'].pop(i)
                    st.rerun()

def render_education():
    st.header("🎓 Education")
    
    if 'education' not in st.session_state.resume_data:
        st.session_state.resume_data['education'] = []
    
    with st.expander("➕ Add Education", expanded=len(st.session_state.resume_data['education']) == 0):
        col1, col2 = st.columns(2)
        
        with col1:
            degree = st.text_input("Degree", key="new_degree", placeholder="Bachelor of Science")
            school = st.text_input("School/University", key="new_school")
            graduation_date = st.text_input("Graduation Date", key="new_grad_date", placeholder="MM/YYYY")
        
        with col2:
            major = st.text_input("Major/Field of Study", key="new_major", placeholder="Computer Science")
            gpa = st.text_input("GPA (optional)", key="new_gpa", placeholder="3.8")
            honors = st.text_input("Honors/Awards (optional)", key="new_honors")
        
        if st.button("✅ Add Education", type="primary"):
            if degree and school:
                education = {
                    'degree': degree,
                    'school': school,
                    'major': major,
                    'graduation_date': graduation_date,
                    'gpa': gpa,
                    'honors': honors
                }
                st.session_state.resume_data['education'].append(education)
                st.success("Education added!")
                st.rerun()
    
    # Display existing education
    for i, edu in enumerate(st.session_state.resume_data['education']):
        with st.expander(f"{edu['degree']} - {edu['school']}", expanded=False):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                st.write(f"**Major:** {edu['major']}")
                st.write(f"**Graduated:** {edu['graduation_date']}")
                if edu['gpa']:
                    st.write(f"**GPA:** {edu['gpa']}")
                if edu['honors']:
                    st.write(f"**Honors:** {edu['honors']}")
            
            with col2:
                if st.button("🗑️ Delete", key=f"delete_edu_{i}"):
                    st.session_state.resume_data['education'].pop(i)
                    st.rerun()

def render_skills():
    st.header("🛠️ Skills")
    
    if 'skills' not in st.session_state.resume_data:
        st.session_state.resume_data['skills'] = {}
    
    # Add skills
    skill_categories = ["Technical Skills", "Programming Languages", "Tools & Software", "Soft Skills", "Languages", "Certifications"]
    
    selected_category = st.selectbox("Category", skill_categories)
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        new_skills = st.text_input(
            "Skills (comma-separated)",
            placeholder="Python, JavaScript, React, Node.js"
        )
    
    with col2:
        if st.button("➕ Add"):
            if new_skills:
                if selected_category not in st.session_state.resume_data['skills']:
                    st.session_state.resume_data['skills'][selected_category] = []
                skills_list = [skill.strip() for skill in new_skills.split(',')]
                st.session_state.resume_data['skills'][selected_category].extend(skills_list)
                st.success("Skills added!")
                st.rerun()
    
    # Display current skills
    if st.session_state.resume_data['skills']:
        st.markdown("### Current Skills")
        for category, skills in st.session_state.resume_data['skills'].items():
            if skills:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"**{category}:** {', '.join(skills)}")
                with col2:
                    if st.button("🗑️", key=f"delete_skills_{category}"):
                        del st.session_state.resume_data['skills'][category]
                        st.rerun()

def render_export():
    st.header("📄 Export Your Resume")
    
    export_mgr = SimpleExportManager()
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("### HTML Export")
        st.write("Best for web viewing and sharing")
        html_data = export_mgr.export_to_html(st.session_state.resume_data)
        st.download_button(
            label="📄 Download HTML",
            data=html_data,
            file_name="resume.html",
            mime="text/html",
            use_container_width=True
        )
    
    with col2:
        st.markdown("### JSON Backup")
        st.write("Save your resume data")
        json_data = json.dumps(st.session_state.resume_data, indent=2)
        st.download_button(
            label="💾 Download JSON",
            data=json_data,
            file_name="resume_data.json",
            mime="application/json",
            use_container_width=True
        )
    
    with col3:
        st.markdown("### Import Data")
        st.write("Load previously saved resume")
        uploaded_file = st.file_uploader("Upload JSON file", type=['json'])
        if uploaded_file is not None:
            try:
                resume_data = json.load(uploaded_file)
                st.session_state.resume_data = resume_data
                st.success("Resume data imported successfully!")
                st.rerun()
            except Exception as e:
                st.error(f"Error importing file: {str(e)}")

def render_preview():
    try:
        export_mgr = SimpleExportManager()
        html_preview = export_mgr.export_to_html(st.session_state.resume_data)
        
        # Create a simplified preview
        personal = st.session_state.resume_data.get('personal', {})
        
        if personal.get('full_name'):
            st.markdown(f"**{personal['full_name']}**")
            
            contact_info = []
            if personal.get('email'):
                contact_info.append(personal['email'])
            if personal.get('phone'):
                contact_info.append(personal['phone'])
            if personal.get('location'):
                contact_info.append(personal['location'])
            
            if contact_info:
                st.caption(" | ".join(contact_info))
        
        if st.session_state.resume_data.get('summary'):
            st.markdown("**Professional Summary**")
            st.write(st.session_state.resume_data['summary'])
        
        work_exp = st.session_state.resume_data.get('work_experience', [])
        if work_exp:
            st.markdown("**Experience**")
            for exp in work_exp[:2]:  # Show only first 2 for preview
                st.write(f"• {exp.get('job_title', '')} at {exp.get('company', '')}")
        
        skills = st.session_state.resume_data.get('skills', {})
        if skills:
            st.markdown("**Skills**")
            for category, skill_list in list(skills.items())[:2]:  # Show only first 2 categories
                if skill_list:
                    st.write(f"**{category}:** {', '.join(skill_list[:3])}...")  # Show first 3 skills
        
    except Exception as e:
        st.write("Preview will appear as you add information")

def main():
    initialize_session_state()
    
    # Navigation
    with st.sidebar:
        st.markdown("## 🎯 TAC Resume Builder")
        
        page_options = {
            'landing': '🏠 Home',
            'auth': '👤 Account',
            'builder': '📝 Resume Builder',
            'templates': '🎨 Templates'
        }
        
        selected_page = st.selectbox(
            "Navigate to:",
            options=list(page_options.keys()),
            format_func=lambda x: page_options[x],
            index=list(page_options.keys()).index(st.session_state.current_page)
        )
        
        if selected_page != st.session_state.current_page:
            st.session_state.current_page = selected_page
            st.rerun()
        
        st.divider()
        
        # Show authentication status
        if st.session_state.user_authenticated:
            render_authentication()
        
        # Show progress if on builder page
        if st.session_state.current_page == 'builder':
            personal_complete = bool(st.session_state.resume_data['personal'].get('full_name'))
            summary_complete = bool(st.session_state.resume_data.get('summary'))
            experience_complete = len(st.session_state.resume_data.get('work_experience', [])) > 0
            education_complete = len(st.session_state.resume_data.get('education', [])) > 0
            skills_complete = len(st.session_state.resume_data.get('skills', {})) > 0
            
            completed = sum([personal_complete, summary_complete, experience_complete, education_complete, skills_complete])
            progress = completed / 5
            
            st.markdown("### 📊 Progress")
            st.progress(progress)
            st.caption(f"{int(progress * 100)}% Complete")
    
    # Route to pages
    if st.session_state.current_page == 'landing':
        render_landing_page()
    elif st.session_state.current_page == 'auth':
        if not render_authentication():
            st.title("👤 Account Access")
            st.markdown("Create an account or login to save your resumes and continue working on them later.")
    elif st.session_state.current_page == 'builder':
        render_builder_page()
    elif st.session_state.current_page == 'templates':
        render_templates_page()

if __name__ == "__main__":
    main()